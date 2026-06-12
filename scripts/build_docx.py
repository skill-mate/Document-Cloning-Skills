#!/usr/bin/env python3
"""
build_docx.py — Render Markdown (with skill directives) into a strictly formatted .docx
using a profile.json (extracted from a sample, or one of the built-in presets).

Hard guarantees over plain python-docx:
  • Cover page is its own section, vertically centered, no page number.
  • TOC page is its own section, with Word TOC field (auto-update on open).
  • Every H1 forces a page break before it (when profile.auto_h1_pagebreak=true).
  • Tables: every row gets w:cantSplit; header row gets w:tblHeader (repeats on each page);
    table layout is fixed.
  • Each run sets BOTH ascii/hAnsi (Western) and eastAsia (CJK) fonts to avoid Word fallback.
  • Page margins, headers/footers, page numbers from profile.

Usage:
    build_docx.py INPUT.md -o OUT.docx [--profile P.json | --preset gov_gongwen|tech_doc]
"""
from __future__ import annotations
import argparse, json, os, re, sys
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
PRESET_DIR = HERE.parent / "presets"

# ---------- profile loading + fallbacks ----------

def _fallback_profile():
    return json.loads((PRESET_DIR / "tech_doc.json").read_text(encoding="utf-8"))

def load_profile(path: str | None, preset: str | None) -> dict:
    if path:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    if preset:
        p = PRESET_DIR / f"{preset}.json"
        if not p.exists():
            print(f"[fallback] preset {preset} not found, using tech_doc", file=sys.stderr)
            p = PRESET_DIR / "tech_doc.json"
        return json.loads(p.read_text(encoding="utf-8"))
    return _fallback_profile()

def pget(prof, dotted, default=None):
    cur = prof
    for k in dotted.split("."):
        if not isinstance(cur, dict) or k not in cur:
            print(f"[fallback] {dotted} -> {default!r}", file=sys.stderr)
            return default
        cur = cur[k]
    return cur

# ---------- low-level XML helpers ----------

def _set_run_fonts(run, eastAsia: str, ascii_: str):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    if ascii_:
        rFonts.set(qn("w:ascii"), ascii_)
        rFonts.set(qn("w:hAnsi"), ascii_)
        rFonts.set(qn("w:cs"), ascii_)
    if eastAsia:
        rFonts.set(qn("w:eastAsia"), eastAsia)
    # also set via python-docx for ascii so Word doesn't drop it
    if ascii_:
        run.font.name = ascii_

def _set_run(run, *, eastAsia, ascii_, size_pt=None, bold=None, italic=None, color_hex=None):
    _set_run_fonts(run, eastAsia, ascii_)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)

def _add_pagebreak(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)

def _para_pagebreak_before(p):
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement("w:pageBreakBefore")
    pPr.append(el)

def _new_section(doc, *, start_type="nextPage"):
    """Add a new section break and return the new section object."""
    sect = doc.add_section(WD_SECTION_START.NEW_PAGE if start_type == "nextPage" else WD_SECTION_START.CONTINUOUS)
    return sect

def _set_section_margins(section, prof):
    section.top_margin = Cm(pget(prof, "page.margin_top_cm", 2.54))
    section.bottom_margin = Cm(pget(prof, "page.margin_bottom_cm", 2.54))
    section.left_margin = Cm(pget(prof, "page.margin_left_cm", 3.17))
    section.right_margin = Cm(pget(prof, "page.margin_right_cm", 3.17))
    section.header_distance = Cm(pget(prof, "page.header_cm", 1.5))
    section.footer_distance = Cm(pget(prof, "page.footer_cm", 1.75))
    section.page_width = Cm(pget(prof, "page.width_cm", 21.0))
    section.page_height = Cm(pget(prof, "page.height_cm", 29.7))

def _set_section_vertical_center(section):
    sectPr = section._sectPr
    vAlign = sectPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        sectPr.append(vAlign)
    vAlign.set(qn("w:val"), "center")

def _clear_section_vertical_align(section):
    sectPr = section._sectPr
    vAlign = sectPr.find(qn("w:vAlign"))
    if vAlign is not None:
        sectPr.remove(vAlign)

def _set_section_different_first_page(section, flag=True):
    section.different_first_page_header_footer = flag

# --- per-level lookup that prefers profile.levels[role], falls back to preset flat keys ---

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

def _level(prof, role):
    """Return the levels[role] sub-dict, or {}."""
    return (prof.get("levels") or {}).get(role, {}) or {}

def _level_get(prof, role, key, default=None):
    return _level(prof, role).get(key, default)

def _resolve_level_fonts(prof, role):
    """Pick (eastAsia, ascii) for a role.
    Priority: profile.levels[role].east/ascii > role-specific font in fonts.* > fonts.default_*"""
    lvl = _level(prof, role)
    if lvl.get("east") or lvl.get("ascii"):
        return (lvl.get("east") or pget(prof, "fonts.default_eastAsia", "宋体"),
                lvl.get("ascii") or pget(prof, "fonts.default_ascii", "Times New Roman"))
    if role.startswith("h"):
        return (pget(prof, "fonts.heading_eastAsia", "黑体"),
                pget(prof, "fonts.heading_ascii", "Arial"))
    if role == "title":
        return (pget(prof, "fonts.title_eastAsia", "黑体"),
                pget(prof, "fonts.title_ascii", "Arial"))
    if role == "table":
        return (pget(prof, "fonts.table_eastAsia", "宋体"),
                pget(prof, "fonts.table_ascii", "Times New Roman"))
    return (pget(prof, "fonts.default_eastAsia", "宋体"),
            pget(prof, "fonts.default_ascii", "Times New Roman"))

def _resolve_level_size(prof, role):
    lvl = _level(prof, role)
    if lvl.get("size_pt"):
        return lvl["size_pt"]
    return pget(prof, f"sizes_pt.{role}", pget(prof, "sizes_pt.body", 12))

def _resolve_level_color(prof, role):
    lvl = _level(prof, role)
    if lvl.get("color"):
        return lvl["color"]
    return pget(prof, f"colors.{role}", pget(prof, "colors.body", "000000"))

def _resolve_level_bold(prof, role):
    lvl = _level(prof, role)
    if "bold" in lvl:
        return lvl["bold"]
    return None  # caller decides default

def _apply_paragraph_format(p, prof, *, role="body"):
    """Apply per-level paragraph formatting. Prefers profile.levels[role] (extracted from
    sample), falls back to preset flat keys (paragraph.*) when not present."""
    pf = p.paragraph_format
    lvl = _level(prof, role)

    # ---- alignment ----
    if "align" in lvl:
        p.alignment = ALIGN_MAP.get(lvl["align"], WD_ALIGN_PARAGRAPH.LEFT)
    elif role == "body":
        p.alignment = ALIGN_MAP.get(pget(prof, "paragraph.body_align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)
    elif role.startswith("h"):
        p.alignment = ALIGN_MAP.get(pget(prof, "paragraph.heading_align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    elif role == "title" or role == "toc_title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- first-line indent ----
    # Priority: levels[role].first_line_indent_pt > (body) profile chars × body size > 0
    if "first_line_indent_pt" in lvl:
        pf.first_line_indent = Pt(lvl["first_line_indent_pt"])
    elif role == "body":
        chars = pget(prof, "paragraph.first_line_indent_chars", 2)
        body_size = pget(prof, "sizes_pt.body", 12)
        pf.first_line_indent = Pt(body_size * chars)

    # ---- left/right indent ----
    if "left_indent_pt" in lvl:
        pf.left_indent = Pt(lvl["left_indent_pt"])
    if "right_indent_pt" in lvl:
        pf.right_indent = Pt(lvl["right_indent_pt"])

    # ---- line spacing ----
    if lvl.get("line_spacing_rule") == "exact" and "line_spacing_pt" in lvl:
        pf.line_spacing = Pt(lvl["line_spacing_pt"])
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    elif lvl.get("line_spacing_rule") == "multiple" and "line_spacing_value" in lvl:
        pf.line_spacing = lvl["line_spacing_value"]
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    else:
        rule = pget(prof, "paragraph.line_spacing_rule", "multiple")
        if rule == "exact":
            pf.line_spacing = Pt(pget(prof, "paragraph.line_spacing_pt", 28))
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        else:
            pf.line_spacing = pget(prof, "paragraph.line_spacing_value", 1.5)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # ---- space before / after ----
    if "space_before_pt" in lvl:
        pf.space_before = Pt(lvl["space_before_pt"])
    elif role == "h1":
        pf.space_before = Pt(pget(prof, "paragraph.h1_space_before_pt", 0))
    elif role == "h2":
        pf.space_before = Pt(pget(prof, "paragraph.h2_space_before_pt", 0))
    else:
        pf.space_before = Pt(pget(prof, "paragraph.space_before_pt", 0))

    if "space_after_pt" in lvl:
        pf.space_after = Pt(lvl["space_after_pt"])
    elif role == "h1":
        pf.space_after = Pt(pget(prof, "paragraph.h1_space_after_pt", 0))
    elif role == "h2":
        pf.space_after = Pt(pget(prof, "paragraph.h2_space_after_pt", 0))
    else:
        pf.space_after = Pt(pget(prof, "paragraph.space_after_pt", 0))

    # ---- keep_with_next / keep_together ----
    if lvl.get("keep_with_next") is not None:
        pf.keep_with_next = bool(lvl["keep_with_next"])
    elif role.startswith("h") or role == "title":
        pf.keep_with_next = True  # sensible default for headings
    if lvl.get("keep_together") is not None:
        pf.keep_together = bool(lvl["keep_together"])

# ---------- Markdown parsing (line-based, just enough for our use) ----------

DIRECTIVE_RE = re.compile(r"^<!--\s*([a-zA-Z][\w-]*)(.*?)-->\s*$", re.S)
INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")

def parse_md(text: str):
    """Return a list of block dicts. Block types:
    - {'kind':'cover', 'fields':{...}}
    - {'kind':'toc', 'depth':int}
    - {'kind':'pagebreak'}
    - {'kind':'newsection'}
    - {'kind':'heading', 'level':1..6, 'text':str}
    - {'kind':'para', 'text':str}
    - {'kind':'quote', 'text':str}
    - {'kind':'list', 'ordered':bool, 'items':[str,...]}
    - {'kind':'table', 'header':[...], 'rows':[[...],...], 'caption':str|None, 'widths':[..]|None}
    - {'kind':'image', 'src':str, 'alt':str, 'caption':str|None}
    - {'kind':'code', 'text':str, 'lang':str}
    """
    lines = text.splitlines()
    i = 0
    blocks = []
    pending_table_caption = None
    pending_table_widths = None
    pending_image_caption = None

    def flush_paragraph(buf):
        if buf:
            blocks.append({"kind": "para", "text": " ".join(s.strip() for s in buf).strip()})
        return []

    para_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # multi-line cover directive
        if stripped.startswith("<!--") and "cover" in stripped and "-->" not in stripped:
            block = [stripped]
            i += 1
            while i < len(lines) and "-->" not in lines[i]:
                block.append(lines[i]); i += 1
            if i < len(lines):
                block.append(lines[i]); i += 1
            joined = "\n".join(block)
            fields = {}
            for m in re.finditer(r"^\s*([a-zA-Z_][\w]*)\s*:\s*(.+?)\s*$", joined, re.M):
                fields[m.group(1)] = m.group(2)
            para_buf = flush_paragraph(para_buf)
            blocks.append({"kind": "cover", "fields": fields})
            continue

        # single-line directive
        m = DIRECTIVE_RE.match(stripped)
        if m:
            name = m.group(1).lower()
            args = m.group(2).strip()
            para_buf = flush_paragraph(para_buf)
            if name == "toc":
                depth_m = re.search(r"depth\s*=\s*(\d+)", args)
                blocks.append({"kind": "toc", "depth": int(depth_m.group(1)) if depth_m else 3})
            elif name == "pagebreak":
                blocks.append({"kind": "pagebreak"})
            elif name == "newsection":
                blocks.append({"kind": "newsection"})
            elif name == "table-caption":
                cap_m = re.match(r":\s*(.+)", args)
                pending_table_caption = cap_m.group(1).strip() if cap_m else args.strip(": ").strip()
            elif name == "image-caption":
                cap_m = re.match(r":\s*(.+)", args)
                pending_image_caption = cap_m.group(1).strip() if cap_m else args.strip(": ").strip()
            elif name == "table-widths":
                wm = re.match(r":\s*(.+)", args)
                raw = wm.group(1) if wm else args.strip(": ").strip()
                try:
                    pending_table_widths = [float(x) for x in re.split(r"[,\s]+", raw) if x]
                except ValueError:
                    pending_table_widths = None
            i += 1; continue

        # blank line
        if stripped == "":
            para_buf = flush_paragraph(para_buf)
            i += 1; continue

        # fenced code
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1; code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i]); i += 1
            i += 1  # skip closing fence
            para_buf = flush_paragraph(para_buf)
            blocks.append({"kind": "code", "text": "\n".join(code_lines), "lang": lang})
            continue

        # heading
        hm = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if hm:
            para_buf = flush_paragraph(para_buf)
            blocks.append({"kind": "heading", "level": len(hm.group(1)), "text": hm.group(2)})
            i += 1; continue

        # blockquote
        if stripped.startswith(">"):
            para_buf = flush_paragraph(para_buf)
            qbuf = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                qbuf.append(lines[i].lstrip()[1:].strip()); i += 1
            blocks.append({"kind": "quote", "text": " ".join(qbuf)})
            continue

        # table (pipe table)
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i+1]):
            para_buf = flush_paragraph(para_buf)
            def split_row(s):
                s = s.strip()
                if s.startswith("|"): s = s[1:]
                if s.endswith("|"): s = s[:-1]
                return [c.strip() for c in s.split("|")]
            header = split_row(lines[i])
            i += 2  # skip separator
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(split_row(lines[i])); i += 1
            blocks.append({
                "kind": "table",
                "header": header,
                "rows": rows,
                "caption": pending_table_caption,
                "widths": pending_table_widths,
            })
            pending_table_caption = None
            pending_table_widths = None
            continue

        # image
        im = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", stripped)
        if im:
            para_buf = flush_paragraph(para_buf)
            blocks.append({"kind": "image", "alt": im.group(1), "src": im.group(2), "caption": pending_image_caption})
            pending_image_caption = None
            i += 1; continue

        # list
        lm = re.match(r"^([-*+])\s+(.+)", stripped)
        om = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if lm or om:
            para_buf = flush_paragraph(para_buf)
            ordered = bool(om)
            items = []
            while i < len(lines):
                s = lines[i].strip()
                m1 = re.match(r"^([-*+])\s+(.+)", s) if not ordered else None
                m2 = re.match(r"^(\d+)\.\s+(.+)", s) if ordered else None
                if (ordered and m2) or (not ordered and m1):
                    items.append((m2 or m1).group(2))
                    i += 1
                elif s == "":
                    break
                else:
                    break
            blocks.append({"kind": "list", "ordered": ordered, "items": items})
            continue

        # paragraph accumulator
        para_buf.append(line)
        i += 1

    flush_paragraph(para_buf)
    return blocks

# ---------- inline run rendering ----------

def add_inline_runs(p, text: str, prof, *, role="body"):
    """Split text by **bold**, *italic*, `code` and add runs with proper fonts.
    Resolves font/size/color from profile.levels[role] first, falling back to flat keys."""
    east, asci = _resolve_level_fonts(prof, role)
    size_pt = _resolve_level_size(prof, role)
    color = _resolve_level_color(prof, role)
    role_bold = _resolve_level_bold(prof, role)  # may be None

    # Tokenize: walk text and find markup spans
    tokens = []
    pos = 0
    pattern = re.compile(r"(\*\*([^*]+?)\*\*|\*([^*\n]+?)\*|`([^`]+?)`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos:m.start()]))
        if m.group(2) is not None:
            tokens.append(("bold", m.group(2)))
        elif m.group(3) is not None:
            tokens.append(("italic", m.group(3)))
        elif m.group(4) is not None:
            tokens.append(("code", m.group(4)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))
    if not tokens:
        tokens = [("plain", text)]

    # Default bold for the role: prefer extracted value, else heading/title=True
    if role_bold is None:
        default_bold = role.startswith("h") or role == "title"
    else:
        default_bold = bool(role_bold)

    for kind, content in tokens:
        if not content:
            continue
        run = p.add_run(content)
        bold = True if kind == "bold" else default_bold
        italic = (kind == "italic")
        if kind == "code":
            _set_run(run, eastAsia=east, ascii_="Consolas", size_pt=size_pt, bold=bold, italic=italic, color_hex=color)
        else:
            _set_run(run, eastAsia=east, ascii_=asci, size_pt=size_pt, bold=bold, italic=italic, color_hex=color)

# ---------- block renderers ----------

def render_cover(doc, fields, prof):
    section = doc.sections[-1]
    _set_section_margins(section, prof)
    _set_section_vertical_center(section)
    _set_section_different_first_page(section, True)
    # blank line for top spacing
    title = fields.get("title", "")
    subtitle = fields.get("subtitle", "")
    org = fields.get("org", "")
    date = fields.get("date", "")

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        _set_run(r,
                 eastAsia=pget(prof, "fonts.title_eastAsia", "黑体"),
                 ascii_=pget(prof, "fonts.title_ascii", "Arial"),
                 size_pt=pget(prof, "cover.title_size_pt", 36),
                 bold=pget(prof, "cover.title_bold", True),
                 color_hex=pget(prof, "cover.title_color", "000000"))
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        r = p.add_run(subtitle)
        _set_run(r,
                 eastAsia=pget(prof, "fonts.title_eastAsia", "黑体"),
                 ascii_=pget(prof, "fonts.title_ascii", "Arial"),
                 size_pt=pget(prof, "cover.subtitle_size_pt", 22),
                 bold=False, color_hex="000000")
    # spacer
    for _ in range(6):
        doc.add_paragraph()
    if org:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(org)
        _set_run(r,
                 eastAsia=pget(prof, "fonts.heading_eastAsia", "黑体"),
                 ascii_=pget(prof, "fonts.heading_ascii", "Arial"),
                 size_pt=pget(prof, "cover.meta_size_pt", 14),
                 bold=True, color_hex="000000")
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(date)
        _set_run(r,
                 eastAsia=pget(prof, "fonts.default_eastAsia", "宋体"),
                 ascii_=pget(prof, "fonts.default_ascii", "Times New Roman"),
                 size_pt=pget(prof, "cover.meta_size_pt", 14))
    # close cover with section break (next page)
    _new_section(doc, start_type="nextPage")
    # remove vertical center on the new section (default is top)
    new_sec = doc.sections[-1]
    _set_section_margins(new_sec, prof)
    _clear_section_vertical_align(new_sec)
    new_sec.different_first_page_header_footer = False

def render_toc(doc, depth, prof):
    # TOC title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(pget(prof, "toc.title", "目  录"))
    _set_run(r,
             eastAsia=pget(prof, "fonts.heading_eastAsia", "黑体"),
             ascii_=pget(prof, "fonts.heading_ascii", "Arial"),
             size_pt=pget(prof, "toc.title_size_pt", 22),
             bold=True, color_hex="000000")
    # TOC field
    p = doc.add_paragraph()
    r = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = f' TOC \\o "1-{depth}" \\h \\z \\u '
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t"); fldChar3.text = "右键此处选择「更新域」以生成目录"
    fldChar4 = OxmlElement("w:fldChar"); fldChar4.set(qn("w:fldCharType"), "end")
    r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2)
    r._r.append(fldChar3); r._r.append(fldChar4)
    # close toc with section break
    _new_section(doc, start_type="nextPage")
    new_sec = doc.sections[-1]
    _set_section_margins(new_sec, prof)
    _clear_section_vertical_align(new_sec)

def render_heading(doc, level, text, prof):
    role = f"h{min(level, 4)}"
    p = doc.add_paragraph()
    style_name = f"Heading {min(level, 9)}"
    try:
        p.style = doc.styles[style_name]
    except KeyError:
        pass
    if level == 1 and pget(prof, "paragraph.auto_h1_pagebreak", True):
        _para_pagebreak_before(p)
    _apply_paragraph_format(p, prof, role=role)
    add_inline_runs(p, text, prof, role=role)

def render_para(doc, text, prof):
    p = doc.add_paragraph()
    _apply_paragraph_format(p, prof, role="body")
    add_inline_runs(p, text, prof, role="body")

def render_quote(doc, text, prof):
    p = doc.add_paragraph()
    _apply_paragraph_format(p, prof, role="body")
    p.paragraph_format.left_indent = Pt(pget(prof, "sizes_pt.body", 12) * 2)
    p.paragraph_format.first_line_indent = Pt(0)
    add_inline_runs(p, text, prof, role="body")

def render_list(doc, items, ordered, prof):
    east = pget(prof, "fonts.default_eastAsia", "宋体")
    asci = pget(prof, "fonts.default_ascii", "Times New Roman")
    size_pt = pget(prof, "sizes_pt.body", 12)
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        _apply_paragraph_format(p, prof, role="body")
        p.paragraph_format.first_line_indent = Pt(0)
        prefix = f"{idx}. " if ordered else "• "
        r = p.add_run(prefix)
        _set_run(r, eastAsia=east, ascii_=asci, size_pt=size_pt)
        add_inline_runs(p, item, prof, role="body")

def render_code(doc, text, prof):
    p = doc.add_paragraph()
    _apply_paragraph_format(p, prof, role="body")
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    _set_run(r, eastAsia=pget(prof, "fonts.default_eastAsia", "宋体"),
             ascii_="Consolas",
             size_pt=pget(prof, "sizes_pt.body", 12) - 1)

def render_image(doc, src, alt, caption, prof):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(src).exists():
        try:
            p.add_run().add_picture(src, width=Cm(14))
        except Exception as e:
            p.add_run(f"[图片：{src} 加载失败 {e}]")
    else:
        p.add_run(f"[图片占位：{alt or src}]")
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        _set_run(r, eastAsia=pget(prof, "fonts.default_eastAsia", "宋体"),
                 ascii_=pget(prof, "fonts.default_ascii", "Times New Roman"),
                 size_pt=pget(prof, "sizes_pt.caption", 10.5),
                 bold=True)

def render_table(doc, header, rows, caption, widths, prof):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        _set_run(r,
                 eastAsia=pget(prof, "fonts.default_eastAsia", "宋体"),
                 ascii_=pget(prof, "fonts.default_ascii", "Times New Roman"),
                 size_pt=pget(prof, "sizes_pt.caption", 10.5),
                 bold=True)

    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = pget(prof, "table.default_style", "Table Grid")
    except KeyError:
        pass

    # fixed layout
    if pget(prof, "table.layout", "fixed") == "fixed":
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)
        table.autofit = False

    # column widths
    if widths and len(widths) == cols:
        # normalize to page text width
        sec = doc.sections[-1]
        usable = sec.page_width - sec.left_margin - sec.right_margin
        total = sum(widths)
        for c in range(cols):
            w = int(usable * widths[c] / total)
            for row in table.rows:
                row.cells[c].width = w

    # header row
    east_h, asci_h = _resolve_level_fonts(prof, "table_header") if (prof.get("levels") or {}).get("table_header") else (pget(prof, "fonts.heading_eastAsia", "黑体"), pget(prof, "fonts.heading_ascii", "Arial"))
    east_b, asci_b = _resolve_level_fonts(prof, "table")
    size_t = _resolve_level_size(prof, "table")
    size_h = _level_get(prof, "table_header", "size_pt", size_t)
    hdr_bold = _level_get(prof, "table_header", "bold", pget(prof, "table.header_bold", True))
    hdr_align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    cell_align = hdr_align_map.get(pget(prof, "table.cell_align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    head_align = hdr_align_map.get(pget(prof, "table.header_align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    valign_map = {"top": WD_ALIGN_VERTICAL.TOP, "center": WD_ALIGN_VERTICAL.CENTER, "bottom": WD_ALIGN_VERTICAL.BOTTOM}
    cell_valign = valign_map.get(pget(prof, "table.cell_valign", "center"), WD_ALIGN_VERTICAL.CENTER)
    hdr_bg = pget(prof, "colors.table_header_bg", None)

    for c, txt in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.vertical_alignment = cell_valign
        para = cell.paragraphs[0]
        para.alignment = head_align
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run(txt)
        _set_run(run, eastAsia=east_h, ascii_=asci_h, size_pt=size_h,
                 bold=hdr_bold,
                 color_hex=pget(prof, "colors.table_header_fg", "000000"))
        if hdr_bg:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hdr_bg)
            tcPr.append(shd)

    # body rows
    for ri, row in enumerate(rows, start=1):
        for c, txt in enumerate(row[:cols]):
            cell = table.rows[ri].cells[c]
            cell.vertical_alignment = cell_valign
            para = cell.paragraphs[0]
            para.alignment = cell_align
            para.paragraph_format.first_line_indent = Pt(0)
            # support \n inside cell
            parts = str(txt).split("<br>") if "<br>" in str(txt) else [str(txt)]
            for pi, ptxt in enumerate(parts):
                if pi > 0:
                    para = cell.add_paragraph()
                    para.alignment = cell_align
                    para.paragraph_format.first_line_indent = Pt(0)
                # render inline (bold/italic) inside cell
                # simplified: just plain run with table fonts
                run = para.add_run(ptxt)
                _set_run(run, eastAsia=east_b, ascii_=asci_b, size_pt=size_t)

    # cantSplit + tblHeader (XML level)
    if pget(prof, "table.cant_split_rows", True):
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement("w:cantSplit")
            trPr.append(cantSplit)
            # min row height
            minh = pget(prof, "table.row_min_height_cm", 0)
            if minh:
                trHeight = OxmlElement("w:trHeight")
                trHeight.set(qn("w:val"), str(int(minh * 567)))  # cm → twips (1 cm = 567)
                trHeight.set(qn("w:hRule"), "atLeast")
                trPr.append(trHeight)
    if pget(prof, "table.repeat_header_row", True):
        trPr = table.rows[0]._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)

# ---------- footer / page number ----------

def setup_section_footer(section, prof, *, hide_pagenum=False):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    align = pget(prof, "header_footer.footer_align", "center")
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    if hide_pagenum or not pget(prof, "header_footer.footer_show_page_number", True):
        return
    fmt = pget(prof, "header_footer.page_number_format", "{n}")
    pre, post = fmt.split("{n}", 1) if "{n}" in fmt else (fmt, "")
    east = pget(prof, "fonts.default_eastAsia", "宋体")
    asci = pget(prof, "fonts.default_ascii", "Times New Roman")
    size = pget(prof, "sizes_pt.footer", 9)

    if pre:
        r = p.add_run(pre)
        _set_run(r, eastAsia=east, ascii_=asci, size_pt=size)
    # PAGE field
    r = p.add_run()
    _set_run(r, eastAsia=east, ascii_=asci, size_pt=size)
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    r._r.append(fc1); r._r.append(it); r._r.append(fc2)
    if post:
        r = p.add_run(post)
        _set_run(r, eastAsia=east, ascii_=asci, size_pt=size)

# ---------- main render ----------

def render(blocks, prof, out_path):
    doc = Document()
    # apply default style (Normal) fonts so anything we forget still uses CJK
    normal = doc.styles["Normal"]
    normal.font.name = pget(prof, "fonts.default_ascii", "Times New Roman")
    normal.font.size = Pt(pget(prof, "sizes_pt.body", 12))
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), pget(prof, "fonts.default_ascii", "Times New Roman"))
    rfonts.set(qn("w:hAnsi"), pget(prof, "fonts.default_ascii", "Times New Roman"))
    rfonts.set(qn("w:eastAsia"), pget(prof, "fonts.default_eastAsia", "宋体"))

    # initial section margins
    _set_section_margins(doc.sections[0], prof)

    # split blocks: cover (if any) -> toc (if any) -> body
    # we'll render in order; cover and toc internally call _new_section.
    has_cover = any(b["kind"] == "cover" for b in blocks)
    has_toc = any(b["kind"] == "toc" for b in blocks)

    cover_section_index = 0
    toc_section_index = None
    body_section_start = None

    sec_idx = 0
    for b in blocks:
        k = b["kind"]
        if k == "cover":
            render_cover(doc, b["fields"], prof)
            sec_idx += 1
            if has_toc:
                toc_section_index = sec_idx
            else:
                body_section_start = sec_idx
        elif k == "toc":
            render_toc(doc, b["depth"], prof)
            sec_idx += 1
            body_section_start = sec_idx
        elif k == "pagebreak":
            _add_pagebreak(doc)
        elif k == "newsection":
            _new_section(doc, start_type="nextPage")
            _set_section_margins(doc.sections[-1], prof)
            _clear_section_vertical_align(doc.sections[-1])
            sec_idx += 1
        elif k == "heading":
            render_heading(doc, b["level"], b["text"], prof)
        elif k == "para":
            render_para(doc, b["text"], prof)
        elif k == "quote":
            render_quote(doc, b["text"], prof)
        elif k == "list":
            render_list(doc, b["items"], b["ordered"], prof)
        elif k == "code":
            render_code(doc, b["text"], prof)
        elif k == "image":
            render_image(doc, b["src"], b["alt"], b["caption"], prof)
        elif k == "table":
            render_table(doc, b["header"], b["rows"], b["caption"], b["widths"], prof)

    # Footers: hide page number on cover (and toc if cover_no_pagenumber); show on body
    cover_hide = pget(prof, "header_footer.cover_no_pagenumber", True)
    for i, sec in enumerate(doc.sections):
        if has_cover and i == 0 and cover_hide:
            setup_section_footer(sec, prof, hide_pagenum=True)
        elif has_toc and ((i == 1 and has_cover) or (i == 0 and not has_cover)) and cover_hide:
            setup_section_footer(sec, prof, hide_pagenum=True)
        else:
            setup_section_footer(sec, prof, hide_pagenum=False)

    doc.save(out_path)

# ---------- CLI ----------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input_md")
    ap.add_argument("-o", "--output", required=True)
    g = ap.add_mutually_exclusive_group()
    g.add_argument("--profile")
    g.add_argument("--preset", choices=["gov_gongwen", "tech_doc"])
    args = ap.parse_args()

    prof = load_profile(args.profile, args.preset)
    md = Path(args.input_md).read_text(encoding="utf-8")
    blocks = parse_md(md)
    render(blocks, prof, args.output)

    # self-check: re-open the doc to verify it isn't corrupt
    try:
        Document(args.output)
        print(f"OK: wrote {args.output} ({Path(args.output).stat().st_size} bytes)")
    except Exception as e:
        print(f"ERROR: produced docx fails to reopen: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
