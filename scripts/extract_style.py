#!/usr/bin/env python3
"""
extract_style.py v1.1 — capture FULL per-level style from a sample .docx
including paragraph-format (alignment, indents, spacing, keep_with_next, etc.)

Output profile.json adds a `levels` block:

  "levels": {
    "title": { "fonts": {...}, "size_pt": .., "bold": .., "color": "..",
               "align": "..", "first_line_indent_pt": .., "left_indent_pt": ..,
               "right_indent_pt": .., "space_before_pt": .., "space_after_pt": ..,
               "line_spacing_rule": "..", "line_spacing_value": .. or "_pt": ..,
               "keep_with_next": true/false, "keep_together": true/false },
    "h1": {...}, "h2": {...}, "h3": {...}, "h4": {...},
    "body": {...},
    "table": {...}      # only fonts/size/bold/align (cell-level)
  }

Bucketing rules:
  • style.name in {"Heading 1", "标题 1", "Title", ...} → that level
  • else if pPr/outlineLvl exists → "h{N+1}"
  • else if style.name suggests body/normal → "body"
  • else → "body"
"""
from __future__ import annotations
import argparse, json, sys, re
from collections import Counter
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
PRESET_DIR = HERE.parent / "presets"

# style name → bucket
STYLE_MAP = {
    "title": "title",
    "标题": "title",
    "heading 1": "h1", "标题 1": "h1",
    "heading 2": "h2", "标题 2": "h2",
    "heading 3": "h3", "标题 3": "h3",
    "heading 4": "h4", "标题 4": "h4",
    "heading 5": "h4", "标题 5": "h4",
    "normal": "body", "正文": "body", "body text": "body",
}

ALIGN_MAP_REVERSE = {0: "left", 1: "center", 2: "right", 3: "justify", 4: "distribute"}


def _bucket_for(p):
    raw = ((p.style.name if p.style else None) or "").strip()
    key = raw.lower()
    if key in STYLE_MAP:
        return STYLE_MAP[key]
    # Chinese names that include a digit
    m = re.match(r"^标题\s*([1-9])$", raw)
    if m:
        n = int(m.group(1))
        return f"h{min(n,4)}"
    # outline level
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        olvl = pPr.find(qn("w:outlineLvl"))
        if olvl is not None:
            try:
                lvl = int(olvl.get(qn("w:val"))) + 1
                if lvl == 0:
                    return "title"
                return f"h{min(lvl,4)}"
            except (TypeError, ValueError):
                pass
    return "body"


def _run_attrs(run, fallback_pPr=None):
    rPr = run._r.find(qn("w:rPr"))
    east = ascii_ = None
    if rPr is not None:
        rf = rPr.find(qn("w:rFonts"))
        if rf is not None:
            east = rf.get(qn("w:eastAsia"))
            ascii_ = rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi"))
    if not ascii_ and run.font.name:
        ascii_ = run.font.name
    size = run.font.size.pt if run.font.size is not None else None
    bold = run.bold
    color = None
    try:
        if run.font.color and run.font.color.rgb is not None:
            color = str(run.font.color.rgb)
    except Exception:
        pass
    return {"east": east, "ascii": ascii_, "size": size, "bold": bold, "color": color}


def _para_attrs(p):
    pf = p.paragraph_format
    pPr = p._p.find(qn("w:pPr"))
    out = {}
    # alignment
    if p.alignment is not None:
        out["align"] = ALIGN_MAP_REVERSE.get(int(p.alignment), None)
    # spacing before/after (Pt)
    if pf.space_before is not None:
        out["space_before_pt"] = round(pf.space_before.pt, 2)
    if pf.space_after is not None:
        out["space_after_pt"] = round(pf.space_after.pt, 2)
    # indents (Pt)
    if pf.first_line_indent is not None:
        out["first_line_indent_pt"] = round(pf.first_line_indent.pt, 2)
    if pf.left_indent is not None:
        out["left_indent_pt"] = round(pf.left_indent.pt, 2)
    if pf.right_indent is not None:
        out["right_indent_pt"] = round(pf.right_indent.pt, 2)
    # line spacing
    if pf.line_spacing is not None:
        rule = str(pf.line_spacing_rule) if pf.line_spacing_rule else ""
        if "EXACT" in rule or "AT_LEAST" in rule:
            try:
                out["line_spacing_rule"] = "exact"
                out["line_spacing_pt"] = round(pf.line_spacing.pt, 2)
            except AttributeError:
                pass
        else:
            out["line_spacing_rule"] = "multiple"
            try:
                out["line_spacing_value"] = round(float(pf.line_spacing), 2)
            except (TypeError, ValueError):
                pass
    # keep_with_next / keep_together
    if pf.keep_with_next is not None:
        out["keep_with_next"] = bool(pf.keep_with_next)
    if pf.keep_together is not None:
        out["keep_together"] = bool(pf.keep_together)
    # widow control
    if pPr is not None:
        if pPr.find(qn("w:pageBreakBefore")) is not None:
            out["page_break_before"] = True
    return out


def _most_common(values):
    """Most common non-None value; returns None if all None/empty."""
    vs = [v for v in values if v is not None and v != ""]
    if not vs:
        return None
    return Counter(vs).most_common(1)[0][0]


def collect(doc: Document):
    levels = {k: {"runs": [], "paras": []} for k in ("title", "h1", "h2", "h3", "h4", "body")}
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        bucket = _bucket_for(p)
        levels[bucket]["paras"].append(_para_attrs(p))
        for r in p.runs:
            if r.text and r.text.strip():
                levels[bucket]["runs"].append(_run_attrs(r))
    # tables
    table_runs = []
    table_header_runs = []
    table_paras = []
    for t in doc.tables:
        for ri, row in enumerate(t.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue
                    table_paras.append(_para_attrs(p))
                    for r in p.runs:
                        if r.text and r.text.strip():
                            attrs = _run_attrs(r)
                            (table_header_runs if ri == 0 else table_runs).append(attrs)
    return levels, {"runs": table_runs, "header_runs": table_header_runs, "paras": table_paras}


def summarize_level(blob):
    runs = blob["runs"]
    paras = blob["paras"]
    out = {}
    if runs:
        for k, attr in (("east", "east"), ("ascii", "ascii"), ("size_pt", "size"), ("color", "color")):
            v = _most_common([r[attr] for r in runs])
            if v is not None:
                out[k] = v
        # Bold: treat None/False as "not bold". Only mark level as bold if the
        # MAJORITY of runs are explicitly True. (Otherwise a single **inline**
        # span would flip the whole level.)
        bold_vals = [bool(r["bold"]) for r in runs]
        if bold_vals and sum(bold_vals) > len(bold_vals) / 2:
            out["bold"] = True
        elif bold_vals and sum(bold_vals) == 0:
            out["bold"] = False
    if paras:
        keys = set()
        for pa in paras:
            keys.update(pa.keys())
        for k in keys:
            v = _most_common([pa.get(k) for pa in paras])
            if v is not None:
                out[k] = v
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("sample")
    ap.add_argument("-o", "--output", required=True)
    args = ap.parse_args()

    doc = Document(args.sample)
    base = json.loads((PRESET_DIR / "tech_doc.json").read_text(encoding="utf-8"))
    levels, table_blob = collect(doc)

    summary = {k: summarize_level(v) for k, v in levels.items()}
    # title fallback to h1 if empty
    if not summary["title"] and summary["h1"]:
        summary["title"] = dict(summary["h1"])

    table_summary = summarize_level({"runs": table_blob["runs"] or table_blob["header_runs"], "paras": table_blob["paras"]})
    table_header_summary = summarize_level({"runs": table_blob["header_runs"], "paras": []})

    # Mirror to flat keys (kept for back-compat with v1.0 build_docx behaviour)
    body = summary["body"]
    h1 = summary["h1"] or summary["title"]
    title = summary["title"] or h1
    table = table_summary
    if body.get("east"):       base["fonts"]["default_eastAsia"] = body["east"]
    if body.get("ascii"):      base["fonts"]["default_ascii"]    = body["ascii"]
    if body.get("size_pt"):    base["sizes_pt"]["body"]          = body["size_pt"]
    if body.get("color"):      base["colors"]["body"]            = body["color"]
    if title.get("east"):      base["fonts"]["title_eastAsia"]   = title["east"]
    if title.get("ascii"):     base["fonts"]["title_ascii"]      = title["ascii"]
    if title.get("size_pt"):   base["sizes_pt"]["title"]         = title["size_pt"]
    if title.get("color"):     base["colors"]["title"]           = title["color"]
    if h1.get("east"):         base["fonts"]["heading_eastAsia"] = h1["east"]
    if h1.get("ascii"):        base["fonts"]["heading_ascii"]    = h1["ascii"]
    for lvl_name in ("h1", "h2", "h3", "h4"):
        s = summary[lvl_name]
        if s.get("size_pt"): base["sizes_pt"][lvl_name] = s["size_pt"]
        if s.get("color"):   base["colors"][lvl_name]   = s["color"]
    if table.get("east"):      base["fonts"]["table_eastAsia"]   = table["east"]
    if table.get("ascii"):     base["fonts"]["table_ascii"]      = table["ascii"]
    if table.get("size_pt"):   base["sizes_pt"]["table"]         = table["size_pt"]

    # Page setup
    sec = doc.sections[0]
    if sec.top_margin:    base["page"]["margin_top_cm"]    = round(sec.top_margin.cm, 2)
    if sec.bottom_margin: base["page"]["margin_bottom_cm"] = round(sec.bottom_margin.cm, 2)
    if sec.left_margin:   base["page"]["margin_left_cm"]   = round(sec.left_margin.cm, 2)
    if sec.right_margin:  base["page"]["margin_right_cm"]  = round(sec.right_margin.cm, 2)
    if sec.page_width:    base["page"]["width_cm"]         = round(sec.page_width.cm, 2)
    if sec.page_height:   base["page"]["height_cm"]        = round(sec.page_height.cm, 2)

    # Body paragraph defaults (back-compat flat keys)
    if "line_spacing_rule" in body:
        base["paragraph"]["line_spacing_rule"] = body["line_spacing_rule"]
        if body["line_spacing_rule"] == "exact" and "line_spacing_pt" in body:
            base["paragraph"]["line_spacing_pt"] = body["line_spacing_pt"]
        elif "line_spacing_value" in body:
            base["paragraph"]["line_spacing_value"] = body["line_spacing_value"]
    if "first_line_indent_pt" in body and body.get("size_pt"):
        base["paragraph"]["first_line_indent_chars"] = round(body["first_line_indent_pt"] / body["size_pt"], 1)
    if "align" in body:
        base["paragraph"]["body_align"] = body["align"]
    if "align" in h1:
        base["paragraph"]["heading_align"] = h1["align"]

    # NEW: full per-level block (build_docx v1.1 prefers this)
    base["levels"] = {k: v for k, v in summary.items() if v}
    base["levels"]["table"] = table_summary
    if table_header_summary:
        base["levels"]["table_header"] = table_header_summary

    Path(args.output).write_text(json.dumps(base, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"OK: wrote profile to {args.output}")
    for lvl in ("title", "h1", "h2", "h3", "h4", "body"):
        s = summary[lvl]
        if s:
            print(f"  [{lvl}] " +
                  f"font={s.get('east','-')}/{s.get('ascii','-')} " +
                  f"size={s.get('size_pt','-')}pt bold={s.get('bold','-')} " +
                  f"align={s.get('align','-')} " +
                  f"space={s.get('space_before_pt','-')}/{s.get('space_after_pt','-')}pt " +
                  f"indent={s.get('first_line_indent_pt','-')}pt " +
                  f"line={s.get('line_spacing_rule','-')}:{s.get('line_spacing_pt') or s.get('line_spacing_value','-')}")
    print(f"  margins (cm): T{base['page']['margin_top_cm']} B{base['page']['margin_bottom_cm']} L{base['page']['margin_left_cm']} R{base['page']['margin_right_cm']}")


if __name__ == "__main__":
    main()
