#!/usr/bin/env python3
"""
build_from_skeleton.py — Use a reference .docx as a SKELETON: preserve ALL of its
styles.xml, numbering.xml, theme/, header*.xml, footer*.xml, settings.xml, and
section properties (page setup, header/footer references). Replace ONLY the body
content with material from Markdown.

This is the most precise way to produce "the same document with new content":
every visual decision the reference made (fonts, sizes, colours, numbering,
header banner, page numbers, page setup) is inherited byte-for-byte.

Usage:
    build_from_skeleton.py INPUT.md --reference REF.docx -o OUT.docx
        [--no-h1-pagebreak]

Strategy (v1.2 — fixes the "直接格式" bug):

  Most real-world Chinese tenders/proposals do NOT customize styles.xml.
  Instead, they apply formatting DIRECTLY to each <w:p> (font, size, indent,
  spacing) on top of plain "Normal" / "Heading N" styles. If we just clear
  the body and call doc.add_paragraph(style="Heading 1"), python-docx writes
  a paragraph that uses the BARE built-in Heading 1, losing all of the
  reference's direct formatting (黑体二号 → Calibri 28pt fallback,
  仿宋_GB2312 → 宋体 default, 首行缩进 → 0, 段距 → 0).

  Fix: BEFORE clearing the body, walk every paragraph and capture per-role
  (cover_title / title / h1..h4 / body / subtitle) the FIRST occurrence's
  full <w:pPr> + <w:rPr> XML. After clearing, when we render new content,
  we deep-copy those captured elements and merge them into the new paragraph
  / run, overriding python-docx's defaults. This preserves direct formatting.
"""
from __future__ import annotations
import argparse, re, sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from build_docx import parse_md  # only the parser is reused


# ============================================================
# Step 1 — capture formatting templates from the reference body
# ============================================================

# Style-name → role lookup. Walk doc.styles to build {style_id: role}.
def _build_role_map(doc):
    role_by_style_id = {}
    for s in doc.styles:
        try:
            sid = s.style_id
            name = (s.name or "").strip()
            nl = name.lower()
            role = None
            if "heading 1" in nl or name == "标题 1": role = "h1"
            elif "heading 2" in nl or name == "标题 2": role = "h2"
            elif "heading 3" in nl or name == "标题 3": role = "h3"
            elif "heading 4" in nl or name == "标题 4": role = "h4"
            elif "heading 5" in nl or name == "标题 5": role = "h4"
            elif nl == "title" or name == "标题": role = "title"
            elif nl == "subtitle" or name == "副标题": role = "subtitle"
            elif "list" in nl or "列表" in name: role = "list"
            elif "quote" in nl or "引用" in name: role = "quote"
            elif "caption" in nl or "题注" in name: role = "caption"
            elif nl == "normal" or name == "正文": role = "body"
            if sid and role:
                role_by_style_id[sid] = role
        except Exception:
            continue
    return role_by_style_id


def _first_real_run_rPr(p_el):
    """Return a deepcopy of the rPr of the first non-empty run in the paragraph."""
    for r_el in p_el.findall(qn("w:r")):
        # consider runs with text or with break/tab as "real"
        t_el = r_el.find(qn("w:t"))
        has_text = t_el is not None and (t_el.text or "").strip()
        if not has_text:
            continue
        rPr = r_el.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr)
        return None
    return None


def _capture_templates(doc):
    """Walk reference body BEFORE clearing it. Capture per-role pPr+rPr templates.

    Returns a dict: role -> {"pPr": elem|None, "rPr": elem|None}.
    Roles include: cover_title, title, h1..h4, subtitle, body, list, quote, caption.
    """
    role_by_id = _build_role_map(doc)
    body = doc.element.body
    templates = {}

    paragraphs = body.findall(qn("w:p"))

    # First non-empty paragraph is treated as "cover_title" if it differs from
    # subsequent body paragraphs (heuristic: it usually has bigger size or different font).
    first_real = None
    for p_el in paragraphs:
        # only consider paragraphs that have actual text
        ts = [t.text for t in p_el.iter(qn("w:t")) if t.text]
        if any((s or "").strip() for s in ts):
            first_real = p_el
            break
    if first_real is not None:
        templates["cover_title"] = {
            "pPr": deepcopy(first_real.find(qn("w:pPr"))) if first_real.find(qn("w:pPr")) is not None else None,
            "rPr": _first_real_run_rPr(first_real),
        }

    # Capture per-role: first paragraph that maps to that role
    for p_el in paragraphs:
        pPr = p_el.find(qn("w:pPr"))
        style_id = None
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_id = pStyle.get(qn("w:val"))
        role = role_by_id.get(style_id, "body")

        if role in templates:
            continue
        # only capture if there's something interesting to capture
        run_rPr = _first_real_run_rPr(p_el)
        if pPr is None and run_rPr is None:
            continue
        templates[role] = {
            "pPr": deepcopy(pPr) if pPr is not None else None,
            "rPr": run_rPr,
        }

    # Ensure 'body' is always present — fall back to first real paragraph if
    # nothing got tagged as body
    if "body" not in templates and "cover_title" in templates:
        templates["body"] = templates["cover_title"]

    return templates


# ============================================================
# Step 2 — apply captured templates onto generated paragraphs/runs
# ============================================================

# Drop these tags when merging templates — they're revision-tracking noise
# and would cause Word to flag false changes.
_PPR_DROP = {qn("w:rsid"), qn("w:rPrChange"), qn("w:pPrChange")}
_RPR_DROP = {qn("w:rsid"), qn("w:rPrChange")}


def _merge_pPr(p, tpl_pPr):
    """Merge children of tpl_pPr into the paragraph's pPr.
    For each tag in the template, replace any existing same-tag child."""
    if tpl_pPr is None:
        return
    cur = p._p.get_or_add_pPr()
    for child in list(tpl_pPr):
        if child.tag in _PPR_DROP:
            continue
        # remove existing same-tag child
        existing = cur.find(child.tag)
        if existing is not None:
            cur.remove(existing)
        cur.append(deepcopy(child))


def _merge_rPr(run, tpl_rPr):
    """Merge children of tpl_rPr into the run's rPr."""
    if tpl_rPr is None or run is None:
        return
    cur = run._r.get_or_add_rPr()
    for child in list(tpl_rPr):
        if child.tag in _RPR_DROP:
            continue
        existing = cur.find(child.tag)
        if existing is not None:
            cur.remove(existing)
        cur.append(deepcopy(child))


# ============================================================
# Step 3 — body bootstrapping (clear contents, keep sectPr)
# ============================================================

def _clear_body_keep_sectpr(doc):
    body = doc.element.body
    final_sectPr_orig = body.find(qn("w:sectPr"))
    final_sectPr = deepcopy(final_sectPr_orig) if final_sectPr_orig is not None else None
    for child in list(body):
        body.remove(child)
    if final_sectPr is not None:
        body.append(final_sectPr)


# ============================================================
# Style discovery (still used for table style binding)
# ============================================================

def _find_style(doc, *candidates):
    by_name = {}
    for s in doc.styles:
        try:
            n = s.name
        except Exception:
            continue
        if n and n not in by_name:
            by_name[n] = s
    for name in candidates:
        if name in by_name:
            return by_name[name]
    return None


def _heading_style(doc, level):
    return _find_style(doc, f"Heading {level}", f"标题 {level}", f"Heading{level}")


def _table_style(doc):
    return _find_style(
        doc,
        "Table Grid", "网格型", "Light Grid", "Light Shading",
        "Medium Shading 1", "Plain Table 1", "Table Grid Light",
    )


def _apply_style(p, style_obj):
    if style_obj is not None:
        p.style = style_obj
    return p


# ============================================================
# Inline run rendering (with template-aware rPr injection)
# ============================================================

INLINE_RE = re.compile(r"(\*\*([^*]+?)\*\*|\*([^*\n]+?)\*|`([^`]+?)`)")


def _add_inline(p, text, tpl_rPr=None):
    """Add runs to paragraph p. Each plain run also receives the captured rPr
    so direct formatting (font/size/color) survives."""
    def add(text_part, *, bold=False, italic=False, code=False):
        r = p.add_run(text_part)
        _merge_rPr(r, tpl_rPr)
        if code:
            r.font.name = "Consolas"
        if bold:   r.bold = True
        if italic: r.italic = True
        return r

    pos = 0
    matched_any = False
    for m in INLINE_RE.finditer(text):
        matched_any = True
        if m.start() > pos:
            add(text[pos:m.start()])
        if m.group(2) is not None:
            add(m.group(2), bold=True)
        elif m.group(3) is not None:
            add(m.group(3), italic=True)
        elif m.group(4) is not None:
            add(m.group(4), code=True)
        pos = m.end()
    if pos < len(text):
        add(text[pos:])
    if not matched_any and not text:
        # add an empty run with rPr so the paragraph still carries the formatting
        add("")


# ============================================================
# Low-level XML helpers
# ============================================================

def _para_pagebreak_before(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


def _add_pagebreak(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    r._r.append(br)


def _new_section(doc):
    return doc.add_section(WD_SECTION.NEW_PAGE)


def _set_section_vertical_center(section):
    sectPr = section._sectPr
    vAlign = sectPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign"); sectPr.append(vAlign)
    vAlign.set(qn("w:val"), "center")


def _clear_section_vertical_align(section):
    sectPr = section._sectPr
    vAlign = sectPr.find(qn("w:vAlign"))
    if vAlign is not None:
        sectPr.remove(vAlign)


# ============================================================
# Block renderers — now template-aware
# ============================================================

def _tpl(templates, role, fallback="body"):
    """Pick the best available template for this role."""
    if role in templates: return templates[role]
    if fallback in templates: return templates[fallback]
    return {"pPr": None, "rPr": None}


def render_cover(doc, fields, templates):
    title    = fields.get("title", "")
    subtitle = fields.get("subtitle", "")
    org      = fields.get("org", "")
    date     = fields.get("date", "")

    _set_section_vertical_center(doc.sections[-1])

    if title:
        p = doc.add_paragraph()
        # Try Title named style first (cheap), then overlay captured cover_title formatting
        title_style = _find_style(doc, "Title", "标题", "Heading 1", "标题 1")
        _apply_style(p, title_style)
        tpl = _tpl(templates, "cover_title", "title")
        _merge_pPr(p, tpl["pPr"])
        # Force center alignment for cover (in case template wasn't centered)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline(p, title, tpl["rPr"])

    if subtitle:
        p = doc.add_paragraph()
        sub_style = _find_style(doc, "Subtitle", "副标题", "Heading 2")
        _apply_style(p, sub_style)
        tpl = _tpl(templates, "subtitle", "h2")
        _merge_pPr(p, tpl["pPr"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline(p, subtitle, tpl["rPr"])

    for _ in range(6):
        doc.add_paragraph()

    if org:
        p = doc.add_paragraph()
        tpl = _tpl(templates, "h2", "body")
        _merge_pPr(p, tpl["pPr"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # bold via inline already; pass tpl rPr for font
        _add_inline(p, f"**{org}**", tpl["rPr"])

    if date:
        p = doc.add_paragraph()
        tpl = _tpl(templates, "body")
        _merge_pPr(p, tpl["pPr"])
        # cover paragraphs shouldn't have first-line indent
        _strip_first_line_indent(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline(p, date, tpl["rPr"])

    _new_section(doc)
    _clear_section_vertical_align(doc.sections[-1])


def _strip_first_line_indent(p):
    """Remove first-line indent from a paragraph (used for cover/heading/cell paragraphs)."""
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None: return
    ind = pPr.find(qn("w:ind"))
    if ind is None: return
    # Drop firstLine attr
    for attr in (qn("w:firstLine"), qn("w:firstLineChars")):
        if ind.get(attr) is not None:
            ind.attrib.pop(attr)


def render_toc(doc, depth, templates):
    title_style = _find_style(doc, "TOC Heading", "Heading 1", "标题 1")
    p = doc.add_paragraph()
    _apply_style(p, title_style)
    tpl = _tpl(templates, "h1", "cover_title")
    _merge_pPr(p, tpl["pPr"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _strip_first_line_indent(p)
    _add_inline(p, "目  录", tpl["rPr"])

    p = doc.add_paragraph()
    r = p.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "1-{depth}" \\h \\z \\u '
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "右键此处选择「更新域」以生成目录"
    fc4 = OxmlElement("w:fldChar"); fc4.set(qn("w:fldCharType"), "end")
    for el in (fc1, instr, fc2, placeholder, fc4):
        r._r.append(el)

    _new_section(doc)
    _clear_section_vertical_align(doc.sections[-1])


def render_heading(doc, level, text, templates, auto_h1_pb=True):
    role = f"h{min(level, 4)}"
    style_obj = _heading_style(doc, min(level, 4))
    p = doc.add_paragraph()
    _apply_style(p, style_obj)
    tpl = _tpl(templates, role, "body")
    _merge_pPr(p, tpl["pPr"])
    _strip_first_line_indent(p)  # headings don't get body's first-line indent
    if level == 1 and auto_h1_pb:
        _para_pagebreak_before(p)
    _add_inline(p, text, tpl["rPr"])


def render_para(doc, text, templates):
    p = doc.add_paragraph()
    tpl = _tpl(templates, "body")
    _merge_pPr(p, tpl["pPr"])
    _add_inline(p, text, tpl["rPr"])


def render_quote(doc, text, templates):
    style = _find_style(doc, "Quote", "引用", "Intense Quote")
    p = doc.add_paragraph()
    _apply_style(p, style)
    tpl = _tpl(templates, "quote", "body")
    _merge_pPr(p, tpl["pPr"])
    _add_inline(p, text, tpl["rPr"])


def render_list(doc, items, ordered, templates):
    if ordered:
        style = _find_style(doc, "List Number", "List Paragraph")
    else:
        style = _find_style(doc, "List Bullet", "List Paragraph")
    tpl = _tpl(templates, "list", "body")
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        _apply_style(p, style)
        _merge_pPr(p, tpl["pPr"])
        _strip_first_line_indent(p)  # lists shouldn't double-indent
        if style is None:
            _add_inline(p, (f"{idx}. " if ordered else "• ") + item, tpl["rPr"])
        else:
            _add_inline(p, item, tpl["rPr"])


def render_code(doc, text, templates):
    style = _find_style(doc, "HTML Preformatted", "Plain Text")
    p = doc.add_paragraph()
    _apply_style(p, style)
    tpl = _tpl(templates, "body")
    _merge_pPr(p, tpl["pPr"])
    _strip_first_line_indent(p)
    r = p.add_run(text)
    _merge_rPr(r, tpl["rPr"])
    r.font.name = "Consolas"


def render_image(doc, src, alt, caption, templates):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _strip_first_line_indent(p)
    if Path(src).exists():
        try:
            p.add_run().add_picture(src, width=Cm(14))
        except Exception as e:
            p.add_run(f"[图片：{src} 加载失败 {e}]")
    else:
        p.add_run(f"[图片占位：{alt or src}]")
    if caption:
        cstyle = _find_style(doc, "Caption", "题注")
        cp = doc.add_paragraph()
        _apply_style(cp, cstyle)
        tpl = _tpl(templates, "caption", "body")
        _merge_pPr(cp, tpl["pPr"])
        _strip_first_line_indent(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline(cp, caption, tpl["rPr"])


def _add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)


def render_table(doc, header, rows, caption, widths, templates):
    if caption:
        cstyle = _find_style(doc, "Caption", "题注")
        cp = doc.add_paragraph()
        _apply_style(cp, cstyle)
        tpl = _tpl(templates, "caption", "body")
        _merge_pPr(cp, tpl["pPr"])
        _strip_first_line_indent(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline(cp, caption, tpl["rPr"])

    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    style_obj = _table_style(doc)
    if style_obj is not None:
        try:
            table.style = style_obj
        except Exception:
            _add_table_borders(table)
    else:
        _add_table_borders(table)

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    table.autofit = False

    if widths and len(widths) == cols:
        sec = doc.sections[-1]
        usable = sec.page_width - sec.left_margin - sec.right_margin
        total = sum(widths)
        for c in range(cols):
            w = int(usable * widths[c] / total)
            for row in table.rows:
                row.cells[c].width = w

    body_tpl = _tpl(templates, "body")

    # header row — bold + centered
    for c, txt in enumerate(header):
        cell = table.rows[0].cells[c]
        para = cell.paragraphs[0]
        _merge_pPr(para, body_tpl["pPr"])
        _strip_first_line_indent(para)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(txt)
        _merge_rPr(r, body_tpl["rPr"])
        r.bold = True
    # body rows
    for ri, row in enumerate(rows, start=1):
        for c, txt in enumerate(row[:cols]):
            cell = table.rows[ri].cells[c]
            para = cell.paragraphs[0]
            _merge_pPr(para, body_tpl["pPr"])
            _strip_first_line_indent(para)
            parts = str(txt).split("<br>") if "<br>" in str(txt) else [str(txt)]
            for pi, ptxt in enumerate(parts):
                if pi > 0:
                    para = cell.add_paragraph()
                    _merge_pPr(para, body_tpl["pPr"])
                    _strip_first_line_indent(para)
                _add_inline(para, ptxt, body_tpl["rPr"])

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


# ============================================================
# Main render driver
# ============================================================

def render(blocks, doc, templates, *, auto_h1_pagebreak=True):
    for b in blocks:
        k = b["kind"]
        if   k == "cover":      render_cover(doc, b["fields"], templates)
        elif k == "toc":        render_toc(doc, b["depth"], templates)
        elif k == "pagebreak":  _add_pagebreak(doc)
        elif k == "newsection":
            _new_section(doc)
            _clear_section_vertical_align(doc.sections[-1])
        elif k == "heading":    render_heading(doc, b["level"], b["text"], templates, auto_h1_pb=auto_h1_pagebreak)
        elif k == "para":       render_para(doc, b["text"], templates)
        elif k == "quote":      render_quote(doc, b["text"], templates)
        elif k == "list":       render_list(doc, b["items"], b["ordered"], templates)
        elif k == "code":       render_code(doc, b["text"], templates)
        elif k == "image":      render_image(doc, b["src"], b["alt"], b["caption"], templates)
        elif k == "table":      render_table(doc, b["header"], b["rows"], b["caption"], b["widths"], templates)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input_md")
    ap.add_argument("--reference", required=True, help="reference .docx (skeleton)")
    ap.add_argument("-o", "--output", required=True)
    ap.add_argument("--no-h1-pagebreak", action="store_true",
                    help="don't auto-page-break before each H1 (default: do)")
    args = ap.parse_args()

    doc = Document(args.reference)
    # ★ Capture templates BEFORE clearing the body
    templates = _capture_templates(doc)
    _clear_body_keep_sectpr(doc)
    blocks = parse_md(Path(args.input_md).read_text(encoding="utf-8"))
    render(blocks, doc, templates, auto_h1_pagebreak=not args.no_h1_pagebreak)
    doc.save(args.output)

    # self-check
    Document(args.output)
    print(f"OK: wrote {args.output} ({Path(args.output).stat().st_size} bytes)")
    print(f"  reference styles: {sum(1 for _ in doc.styles)}")
    print(f"  captured templates: {sorted(templates.keys())}")
    used_h = []
    for n in (1, 2, 3, 4):
        s = _heading_style(doc, n)
        if s is not None:
            used_h.append(s.name)
    print(f"  heading styles bound: {', '.join(used_h) or '(none)'}")
    ts = _table_style(doc)
    print(f"  table style bound: {ts.name if ts else '(default Table Grid)'}")


if __name__ == "__main__":
    main()
